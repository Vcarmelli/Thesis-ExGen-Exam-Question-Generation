from flask import Blueprint, send_file, session
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from bs4 import BeautifulSoup

key_export = Blueprint('key_export', __name__)

@key_export.route('/export-keynotes-pdf')
def export_keynotes_pdf():
    # Get generated keynotes from session
    keynotes_html = session.get('generated_keynotes', '')
    
    # Parse HTML content
    soup = BeautifulSoup(keynotes_html, 'html.parser')
    keynotes_text = soup.get_text('\n', strip=True)
    
    # Create PDF object
    pdf = FPDF()
    pdf.add_page()
    
    # Add title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Generated Keynotes", ln=True, align='C')
    pdf.ln(10)
    
    # Add content
    pdf.set_font("Arial", size=12)
    lines = keynotes_text.split('\n')
    for line in lines:
        if line.strip():  # Only process non-empty lines
            pdf.multi_cell(0, 10, line.strip())
            pdf.ln(5)
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
        pdf.output(tmp.name)
        
        response = send_file(
            tmp.name,
            download_name='keynotes.pdf',
            as_attachment=True,
            mimetype='application/pdf'
        )
        response._file_to_cleanup = tmp.name
        return response

@key_export.route('/export-keynotes-docx')
def export_keynotes_docx():
    # Get generated keynotes from session
    keynotes_html = session.get('generated_keynotes', '')
    
    # Parse HTML content
    soup = BeautifulSoup(keynotes_html, 'html.parser')
    keynotes_text = soup.get_text('\n', strip=True)
    
    # Create document object
    doc = Document()
    
    # Add title
    title = doc.add_heading('Generated Keynotes', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Add some space after title
    
    # Add content
    lines = keynotes_text.split('\n')
    for line in lines:
        if line.strip():  # Only process non-empty lines
            paragraph = doc.add_paragraph()
            paragraph.add_run(line.strip())
            paragraph.space_after = Pt(12)  # Add space after each paragraph
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        
        response = send_file(
            tmp.name,
            download_name='keynotes.docx',
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response._file_to_cleanup = tmp.name
        return response

# Cleanup function to delete temporary files
def cleanup_temp_file(filepath):
    try:
        os.unlink(filepath)
    except Exception as e:
        print(f"Error deleting temporary file: {e}")