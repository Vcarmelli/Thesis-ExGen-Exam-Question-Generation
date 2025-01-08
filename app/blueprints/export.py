from flask import Blueprint, send_file, session
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import io
import tempfile
import os

export = Blueprint('export', __name__)

@export.route('/export-pdf')
def export_pdf():
    # Get generated questions from session
    generated_questions = session.get('generated_questions', [])
    
    # Create PDF object
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Add content to PDF
    for question_set in generated_questions:
        # Add question type header
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, f"{question_set['type']} Questions:", ln=True)
        pdf.set_font("Arial", size=12)
        
        # Add questions
        for i, question in enumerate(question_set['questions'], 1):
            # Question
            pdf.multi_cell(0, 10, f"Question {i}: {question['question']}")
            
            # Options (if they exist)
            if question.get('options'):
                for letter, option in question['options'].items():
                    pdf_text = f"{letter}) {option}"
                    pdf.multi_cell(0, 10, pdf_text)
            
            # Answer
            pdf.set_font("Arial", 'B', 12)
            pdf.multi_cell(0, 10, f"Answer: {question['answer']}")
            pdf.set_font("Arial", size=12)
            pdf.ln(5)
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
        pdf.output(tmp.name)
        
        return send_file(
            tmp.name,
            download_name='generated_questions.pdf',
            as_attachment=True,
            mimetype='application/pdf'
        )

@export.route('/export-docx')
def export_docx():
    # Get generated questions from session
    generated_questions = session.get('generated_questions', [])
    
    # Create document object
    doc = Document()
    
    # Add content to document
    for question_set in generated_questions:
        # Add question type header
        doc.add_heading(f"{question_set['type']} Questions:", level=1)
        
        # Add questions
        for i, question in enumerate(question_set['questions'], 1):
            # Question
            doc.add_paragraph(f"Question {i}: {question['question']}")
            
            # Options (if they exist)
            if question.get('options'):
                options_para = doc.add_paragraph()
                for letter, option in question['options'].items():
                    options_para.add_run(f"{letter}) {option}\n")
            
            # Answer
            answer_para = doc.add_paragraph()
            answer_para.add_run(f"Answer: {question['answer']}").bold = True
            doc.add_paragraph()  # Add spacing between questions
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        
        return send_file(
            tmp.name,
            download_name='generated_questions.docx',
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

# Cleanup function to delete temporary files
def cleanup_temp_file(filepath):
    try:
        os.unlink(filepath)
    except Exception as e:
        print(f"Error deleting temporary file: {e}")