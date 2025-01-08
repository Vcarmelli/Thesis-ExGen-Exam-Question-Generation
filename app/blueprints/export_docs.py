from flask import Blueprint, session, send_file
import io
from docx import Document
from bs4 import BeautifulSoup

export_docs = Blueprint('export_docs', __name__)

@export_docs.route('/export_docs')
def export_to_docx():
    # Get the keynotes from the session
    generated_keynotes = session.get('generated_keynotes', 'No keynotes available.')

    # Parse the HTML using BeautifulSoup
    soup = BeautifulSoup(generated_keynotes, "html.parser")

    # Create a DOCX document
    doc = Document()
    doc.add_heading('Generated Keynotes', level=1)

    # Iterate through sections in the parsed HTML
    for section in soup.find_all(['h3', 'p', 'div', 'span']):
        if section.name == 'h3':  # Add section titles
            doc.add_heading(section.text.strip(), level=2)
        elif section.name == 'p':  # Add paragraph content
            doc.add_paragraph(section.text.strip())
        elif section.name == 'div':  # Handle div content
            div_class = section.get('class', [])
            if 'numbered-item' in div_class:  # Numbered list item
                doc.add_paragraph(section.text.strip(), style='List Number')
            elif 'bullet-item' in div_class:  # Bullet list item
                doc.add_paragraph(section.text.strip(), style='List Bullet')
        elif section.name == 'span':  # Handle standalone spans
            if section.text.strip():  # Only add non-empty text
                doc.add_paragraph(section.text.strip())

    # Save the document to a BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Send the file to the user
    return send_file(
        doc_io,
        as_attachment=True,
        download_name='keynotes.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
